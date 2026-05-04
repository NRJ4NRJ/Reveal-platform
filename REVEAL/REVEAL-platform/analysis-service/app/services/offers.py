from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph


TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates"


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    full_text = paragraph.text
    if not full_text:
        return
    if paragraph.runs:
        changed = False
        for run in paragraph.runs:
            updated_run_text = run.text
            for source, target in replacements.items():
                updated_run_text = updated_run_text.replace(source, target)
            if updated_run_text != run.text:
                run.text = updated_run_text
                changed = True
        if changed:
            return
    updated_text = full_text
    for source, target in replacements.items():
        updated_text = updated_text.replace(source, target)
    if updated_text != full_text:
        paragraph.text = updated_text


def _replace_in_cell(cell, replacements: dict[str, str]) -> None:
    for paragraph in cell.paragraphs:
        _replace_in_paragraph(paragraph, replacements)


def _set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = text


def _set_paragraph_segments(paragraph, segments: list[dict[str, Any]]) -> None:
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.text = ""
    first_run = paragraph.runs[0]
    first_segment = segments[0] if segments else {"text": ""}
    first_run.text = str(first_segment.get("text", ""))
    if "bold" in first_segment:
        first_run.bold = bool(first_segment["bold"])
    for segment in segments[1:]:
        run = paragraph.add_run(str(segment.get("text", "")))
        if "bold" in segment:
            run.bold = bool(segment["bold"])


def _set_cell_text(cell, text: str) -> None:
    if cell.paragraphs:
        _set_paragraph_text(cell.paragraphs[0], text)
        for paragraph in cell.paragraphs[1:]:
            _set_paragraph_text(paragraph, "")
    else:
        cell.text = text


def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None  # type: ignore[attr-defined]


def _remove_table_column(table, col_index: int) -> None:
    for row in table.rows:
        if len(row.cells) > col_index:
            cell = row.cells[col_index]
            tc = cell._tc
            parent = tc.getparent()
            if parent is not None:
                parent.remove(tc)


def build_nbo_docx(
    *,
    template_name: str,
    replacements: dict[str, str],
    paragraph_updates: list[dict[str, Any]] | None = None,
    table_updates: list[dict[str, Any]] | None = None,
    paragraph_deletions: list[int] | None = None,
    table_column_removals: list[dict[str, int]] | None = None,
) -> bytes:
    template_path = TEMPLATE_DIR / template_name
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    document = Document(template_path)

    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_cell(cell, replacements)

    for update in paragraph_updates or []:
        paragraph_index = int(update["index"])
        if update.get("segments"):
            _set_paragraph_segments(document.paragraphs[paragraph_index], list(update["segments"]))
        else:
            _set_paragraph_text(document.paragraphs[paragraph_index], str(update.get("text", "")))

    for update in table_updates or []:
        table_index = int(update["table"])
        row_index = int(update["row"])
        col_index = int(update["col"])
        _set_cell_text(document.tables[table_index].rows[row_index].cells[col_index], str(update.get("text", "")))

    for update in table_column_removals or []:
        _remove_table_column(document.tables[int(update["table"])], int(update["col"]))

    for paragraph_index in sorted((paragraph_deletions or []), reverse=True):
        if 0 <= paragraph_index < len(document.paragraphs):
            _delete_paragraph(document.paragraphs[paragraph_index])

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
