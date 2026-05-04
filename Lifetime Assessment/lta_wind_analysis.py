"""
Long-Term Wind Distribution Analysis — Baudignécourt Wind Farm
==============================================================
REpower MM92 × 6  |  2.05 MW each  |  Meuse (55), Grand Est, France

Methodology:
  1. Load all 6 SCADA turbines; clean frozen sensor data and outliers
  2. Select turbine with highest mean wind speed (conservative lifetime approach)
  3. Correlate cleaned data with 8 reanalysis points (4 ERA5 @ 100 m + 4 MERRA-2 @ 50 m)
  4. Select best-correlated point → apply MCP to full LTA period
  5. Compute Weibull distribution, wind rose, annual average wind speeds
  6. Produce Word document report
"""

import io
import os
import sys
import warnings
from datetime import datetime
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.lines import Line2D
from scipy.stats import weibull_min
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
warnings.filterwarnings("ignore")
REPORT_RUN_DATE = datetime.now().strftime("%d/%m/%Y")

try:
    import contextily as ctx
    CTX_OK = True
except ImportError:
    CTX_OK = False

sys.stdout.reconfigure(encoding="utf-8")

# ─────────────────────────────────────────────────────────────────────────────
# PATHS
# ─────────────────────────────────────────────────────────────────────────────
BASE         = os.path.dirname(__file__)
SCADA_DIR    = os.path.join(BASE, "9 Wind SCADA data")
ERA5_BP_DIR  = os.path.join(BASE, "0 Donnees ERA & MERRA", "ERA",   "processed", "by_point", "csv")
MERRA_BP_DIR = os.path.join(BASE, "0 Donnees ERA & MERRA", "MERRA", "processed", "by_point", "csv")
OUT_DOCX     = os.path.join(BASE, "Baudignecourt_LTA_Wind_Analysis.docx")

# ─────────────────────────────────────────────────────────────────────────────
# SITE & TURBINE CONFIGURATION
# ─────────────────────────────────────────────────────────────────────────────
SITE_NAME     = "Baudignécourt"
TURBINE_MODEL = "REpower MM92"
RATED_POWER   = 2050      # kW
ROTOR_DIAM    = 92.5      # m

# Individual turbine GPS (from project database 1&3.xlsx)
# Keys E1–E6 correspond to WTG 1–6
TURBINE_COORDS = {
    "E1": (48.565320, 5.443333),
    "E2": (48.565500, 5.443900),
    "E3": (48.565600, 5.436600),
    "E4": (48.564600, 5.431700),
    "E5": (48.563700, 5.425600),
    "E6": (48.563100, 5.421100),
}
TURBINE_IDS = ["E1", "E2", "E3", "E4", "E5", "E6"]

# Site centroid
SITE_LAT = np.mean([v[0] for v in TURBINE_COORDS.values()])
SITE_LON = np.mean([v[1] for v in TURBINE_COORDS.values()])

# ─────────────────────────────────────────────────────────────────────────────
# WIND & REANALYSIS CONFIGURATION
# ─────────────────────────────────────────────────────────────────────────────
# Hub height — confirmed 100 m from project database (1&3.xlsx, col Hub height)
HUB_HEIGHT   = 100    # m  (CONFIRMED from 1&3.xlsx)

ERA5_HEIGHT  = 100    # m  (u100/v100 component height)
MERRA_HEIGHT = 50     # m  (U50M/V50M component height)
ALPHA_SHEAR  = 0.20   # Hellmann wind shear exponent (power law)

# Height correction factors (reanalysis level → hub height)
ERA5_CORR  = (HUB_HEIGHT / ERA5_HEIGHT)  ** ALPHA_SHEAR
MERRA_CORR = (HUB_HEIGHT / MERRA_HEIGHT) ** ALPHA_SHEAR

# LTA reference period
LTA_START = "2005-01-01"
LTA_END   = "2025-12-31"

# 8 reanalysis points to test: 4 ERA5 (0.25° grid) + 4 MERRA-2 (~0.5° grid)
# All 4 ERA5 points surrounding site (48.5646°N, 5.4337°E)
ERA5_POINTS  = [(48.50, 5.25), (48.50, 5.50), (48.75, 5.25), (48.75, 5.50)]
# All 4 MERRA-2 points
MERRA_POINTS = [(48.50, 5.000), (48.50, 5.625), (49.00, 5.000), (49.00, 5.625)]

# Data cleaning
FROZEN_MIN_RECORDS = 6   # identical non-zero values for ≥ 6 × 10-min records → frozen

# REpower MM92 reference power curve
# Source: REpower Systems AG official power curve document (air density 1.225 kg/m³)
# Used for curtailment / derating detection — avoids circular-reference problem of
# data-derived thresholds (e.g. P95), which are pulled down by frequent curtailment.
MM92_WS_REF = np.array([
    0.0,  2.5,  3.0,  3.5,  4.0,  4.5,  5.0,  5.5,
    6.0,  6.5,  7.0,  7.5,  8.0,  8.5,  9.0,  9.5,
   10.0, 10.5, 11.0, 12.0, 24.5, 25.0, 99.0
])
MM92_PW_REF = np.array([
      0,    0,   20,   57,   94,  150,  205,  298,
    391,  518,  645,  812,  979, 1177, 1375, 1750,
   1950, 2000, 2000, 2000, 2000,    0,    0
])
# Flag if measured power < this fraction of the reference curve.
# 0.75 (75 %) provides tolerance for normal scatter and the nacelle-wake offset
# (nacelle anemometer reads ~5–8 % below free-stream, shifting the apparent curve).
MM92_CURT_THRESHOLD = 0.75
# Minimum wind speed for curtailment detection.
# Below this value, 10-min power averages show wide natural scatter (turbine
# start-up/shutdown, turbulence-averaging, yaw transients) indistinguishable from
# curtailment.  The manufacturer curve assumes steady-state conditions rarely met
# below ~6 m/s.  At 6 m/s the reference is 391 kW → threshold ~293 kW, large
# enough to be unambiguously anomalous.
MM92_CURT_WS_MIN = 6.0

# Zero wind-speed inconsistency threshold for cleaning:
# remove ws == 0 m/s only when turbine is clearly producing power.
ZERO_WS_PMIN_KW = 0.01 * RATED_POWER  # 1% rated power

# ─────────────────────────────────────────────────────────────────────────────
# COLOUR PALETTE (8.2 house style)
# ─────────────────────────────────────────────────────────────────────────────
BLUE   = "#003366"
LB     = "#0066CC"
ORANGE = "#CC4400"

# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color="003366"):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before      = Pt(10)
    p.paragraph_format.space_after       = Pt(4)
    p.paragraph_format.page_break_before = False
    r = p.add_run(text)
    r.bold           = True
    r.font.size      = Pt(14 if level == 1 else 12 if level == 2 else 11)
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.name      = "Open Sans"
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.size = Pt(10)
        r.font.name = "Open Sans"
    return p

def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        r.font.size = Pt(9)
        r.italic    = True
        r.font.name = "Open Sans"
    return p

def table_header(tbl, headers, bg="003366"):
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        c.text = h
        set_cell_bg(c, bg)
        c.paragraphs[0].runs[0].bold           = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        c.paragraphs[0].runs[0].font.size      = Pt(9)
        c.paragraphs[0].runs[0].font.name      = "Open Sans"

def haversine_km(lat1, lon1, lat2, lon2):
    R    = 6371.0
    phi1 = np.radians(lat1); phi2 = np.radians(lat2)
    dphi = np.radians(lat2 - lat1)
    dlam = np.radians(lon2 - lon1)
    a    = np.sin(dphi/2)**2 + np.cos(phi1)*np.cos(phi2)*np.sin(dlam/2)**2
    return 2 * R * np.arcsin(np.sqrt(a))

# ─────────────────────────────────────────────────────────────────────────────
# 1. LOAD & CLEAN SCADA DATA
# ─────────────────────────────────────────────────────────────────────────────
print("="*60)
print("STEP 1 — Loading and cleaning SCADA data")
print("="*60)
print(f"  Turbine model  : {TURBINE_MODEL}  ({RATED_POWER} kW, Ø{ROTOR_DIAM} m)")
print(f"  Hub height     : {HUB_HEIGHT} m")
print(f"  Frozen threshold: ≥ {FROZEN_MIN_RECORDS} identical non-zero consecutive records (≥ 60 min)")
print(f"  Zero WS @ power threshold: ws=0 and power>{ZERO_WS_PMIN_KW:.1f} kW")
print()

def clean_turbine(turbine_id):
    fp = os.path.join(SCADA_DIR, turbine_id, f"{turbine_id}_SCADA_all.csv")
    df = pd.read_csv(fp, encoding="utf-8-sig", parse_dates=["timestamp"],
                     low_memory=False)
    df = df.rename(columns={"timestamp": "ts"}).set_index("ts")
    df = df.sort_index()
    df["ws"] = pd.to_numeric(df["windSpeedAvg"],  errors="coerce")
    df["wd"] = pd.to_numeric(df["windDirection"], errors="coerce")
    df["pw"] = pd.to_numeric(df["powerAvg"],      errors="coerce")
    ws_raw  = df["ws"].copy()   # preserve raw wind speed before any cleaning
    n_total = len(df)

    # ── Step A: physical outliers ──────────────────────────────────────────
    bad_ws = (df["ws"] < 0) | (df["ws"] > 50)
    bad_wd = (df["wd"] < 0) | (df["wd"] > 360)
    df.loc[bad_ws, "ws"] = np.nan
    df.loc[bad_wd, "wd"] = np.nan
    n_outlier = int(bad_ws.sum())

    # ── Step B: frozen sensor — non-zero identical value for ≥ threshold ──
    ws_s    = df["ws"].copy()
    changed = ws_s.ne(ws_s.shift()).fillna(True)
    grp_id  = changed.cumsum()
    run_len = grp_id.map(grp_id.value_counts())
    frozen  = (run_len >= FROZEN_MIN_RECORDS) & (ws_s > 0.1) & ws_s.notna()
    df.loc[frozen, "ws"] = np.nan
    n_frozen = int(frozen.sum())

    # ── Step C: inconsistent zero wind speed while producing power ──────────────
    zero_ws_power = (df["ws"] == 0) & df["pw"].notna() & (df["pw"] > ZERO_WS_PMIN_KW)
    df.loc[zero_ws_power, "ws"] = np.nan
    n_zero_ws_power = int(zero_ws_power.sum())

    # ── Step D: curtailment / derating — manufacturer reference curve ──────────
    # Flag records where power < MM92_CURT_THRESHOLD × reference curve.
    # Wind speed is RETAINED for MCP — nacelle anemometer reads independently
    # of turbine operating state.
    curt = pd.Series(False, index=df.index)
    if df["pw"].notna().any():
        ws_v      = ws_raw.fillna(-1.0).values
        ref_pw    = pd.Series(np.interp(ws_v, MM92_WS_REF, MM92_PW_REF), index=df.index)
        in_range  = (ws_raw >= MM92_CURT_WS_MIN) & (ws_raw < 25.0) & df["pw"].notna() & (ref_pw > 0)
        curt      = in_range & (df["pw"] < MM92_CURT_THRESHOLD * ref_pw)
        n_curt    = int(curt.sum())
    else:
        n_curt = 0

    n_clean = int(df["ws"].notna().sum())
    stats = {
        "turbine":       turbine_id,
        "n_total":       n_total,
        "n_outlier":     n_outlier,
        "n_frozen":    n_frozen,
        "n_zero_ws_power": n_zero_ws_power,
        "n_curtailed": n_curt,
        "n_clean":     n_clean,
        "avail_pct":     100 * n_clean / n_total,
        "mean_ws":       round(float(df["ws"].mean()), 3),
    }
    flagged            = bad_ws | frozen | zero_ws_power
    df["ws_raw"]       = ws_raw
    df["flagged"]      = flagged
    df["curtailed"] = curt
    return df[["ws", "wd", "pw", "ws_raw", "flagged", "curtailed"]], stats

turbine_data  = {}
turbine_stats = []
for tid in TURBINE_IDS:
    df_t, st = clean_turbine(tid)
    turbine_data[tid] = df_t
    turbine_stats.append(st)
    print(f"  {tid}: mean={st['mean_ws']:.2f} m/s  avail={st['avail_pct']:.1f}%  "
          f"frozen={st['n_frozen']:,}  zero@power={st['n_zero_ws_power']:,}  "
          f"curt={st['n_curtailed']:,}")

stats_df = pd.DataFrame(turbine_stats).set_index("turbine")

# ── Select turbine with highest mean wind speed ────────────────────────────
best_turbine = str(stats_df["mean_ws"].idxmax())
best_ws_mean = float(stats_df.loc[best_turbine, "mean_ws"])
print(f"\n  Selected turbine: {best_turbine}  "
      f"(mean WS = {best_ws_mean:.2f} m/s — conservative for lifetime)")

df_best   = turbine_data[best_turbine][["ws"]].copy()
df_best.columns = ["ws_turb"]
df_best_h = df_best.resample("1h").mean()
df_best_h = df_best_h[df_best_h["ws_turb"].notna()]
print(f"  {best_turbine} hourly valid: {len(df_best_h):,} rows  "
      f"({df_best_h.index.min().date()} – {df_best_h.index.max().date()})")

# ─────────────────────────────────────────────────────────────────────────────
# 2. LOAD ALL 8 REANALYSIS POINTS
# ─────────────────────────────────────────────────────────────────────────────
print()
print("="*60)
print("STEP 2 — Loading 8 reanalysis grid points")
print("="*60)
print(f"  ERA5  height correction : {ERA5_HEIGHT} m → {HUB_HEIGHT} m  "
      f"(factor = {ERA5_CORR:.4f})")
print(f"  MERRA2 height correction: {MERRA_HEIGHT} m → {HUB_HEIGHT} m  "
      f"(factor = {MERRA_CORR:.4f})")
print()

def _ws_wd_from_uv(u, v):
    ws = np.sqrt(u**2 + v**2)
    wd = (np.degrees(np.arctan2(-u, -v)) % 360)
    return ws, wd

def load_era5_point(lat, lon):
    lon_s = f"{lon:.2f}".replace(".", "p")
    lat_s = f"{lat:.2f}".replace(".", "p")
    fname = f"Baudign\u00e9court_point_({lon_s}_{lat_s}).csv"
    fp    = os.path.join(ERA5_BP_DIR, fname)
    df    = pd.read_csv(fp, parse_dates=["time"], low_memory=False,
                        usecols=["time", "u100", "v100"])
    df    = df.dropna(subset=["u100", "v100"])
    ws, wd = _ws_wd_from_uv(df["u100"].values, df["v100"].values)
    df["ws"] = ws * ERA5_CORR
    df["wd"] = wd
    df = df.set_index("time")[["ws", "wd"]].resample("1h").mean()
    return df.loc[LTA_START:LTA_END]

def load_merra_point(lat, lon):
    lon_s = f"{lon:.3f}".replace(".", "p")
    lat_s = f"{lat:.3f}".replace(".", "p")
    fname = f"Baudign\u00e9court_point_({lon_s}_{lat_s}).csv"
    fp    = os.path.join(MERRA_BP_DIR, fname)
    df    = pd.read_csv(fp, parse_dates=["time"], low_memory=False,
                        usecols=["time", "U50M", "V50M"])
    df    = df.dropna(subset=["U50M", "V50M"])
    ws, wd = _ws_wd_from_uv(df["U50M"].values, df["V50M"].values)
    df["ws"] = ws * MERRA_CORR
    df["wd"] = wd
    df = df.set_index("time")[["ws", "wd"]].resample("1h").mean()
    return df.loc[LTA_START:LTA_END]

reanalysis_pts = {}   # (dataset, lat, lon) → DataFrame
for (lat, lon) in ERA5_POINTS:
    key = ("ERA5", lat, lon)
    reanalysis_pts[key] = load_era5_point(lat, lon)
    dist = haversine_km(SITE_LAT, SITE_LON, lat, lon)
    print(f"  ERA5   ({lat:.2f}°N, {lon:.2f}°E)  dist={dist:.1f} km  "
          f"rows={len(reanalysis_pts[key]):,}")

for (lat, lon) in MERRA_POINTS:
    key = ("MERRA-2", lat, lon)
    reanalysis_pts[key] = load_merra_point(lat, lon)
    dist = haversine_km(SITE_LAT, SITE_LON, lat, lon)
    print(f"  MERRA-2 ({lat:.2f}°N, {lon:.3f}°E)  dist={dist:.1f} km  "
          f"rows={len(reanalysis_pts[key]):,}")

# ─────────────────────────────────────────────────────────────────────────────
# 3. CORRELATION — TURBINE Ebest vs ALL 8 POINTS
# ─────────────────────────────────────────────────────────────────────────────
print()
print("="*60)
print(f"STEP 3 — Correlating {best_turbine} with all 8 reanalysis points")
print("="*60)

corr_rows = []
overlap_cache = {}
for (dataset, lat, lon), ref_df in reanalysis_pts.items():
    merged = df_best_h.join(ref_df[["ws"]], how="inner").dropna()
    merged.columns = ["ws_turb", "ws_ref"]
    n_ov = len(merged)
    if n_ov < 100:
        r, a, b = np.nan, np.nan, np.nan
    else:
        r    = float(np.corrcoef(merged["ws_ref"], merged["ws_turb"])[0, 1])
        a, b = np.polyfit(merged["ws_ref"], merged["ws_turb"], 1)
    dist = haversine_km(SITE_LAT, SITE_LON, lat, lon)
    corr_rows.append(dict(dataset=dataset, lat=lat, lon=lon,
                          dist_km=dist, n_overlap=n_ov, R=r, slope=a, intercept=b))
    overlap_cache[(dataset, lat, lon)] = merged
    print(f"  {dataset:7s} ({lat:.2f},{lon:.3f})  "
          f"dist={dist:5.1f} km  n={n_ov:,}  R={r:.4f}")

corr_df = pd.DataFrame(corr_rows)

# Select best point
best_idx     = int(corr_df["R"].idxmax())
best_row     = corr_df.iloc[best_idx]
best_dataset = str(best_row["dataset"])
best_lat     = float(best_row["lat"])
best_lon     = float(best_row["lon"])
best_R       = float(best_row["R"])
best_slope   = float(best_row["slope"])
best_intcpt  = float(best_row["intercept"])
_sign_intcpt = "+ " if best_intcpt >= 0 else "- "

print(f"\n  --> Best point: {best_dataset} ({best_lat}°N, {best_lon}°E)  "
      f"R = {best_R:.4f}")
print(f"      SCADA = {best_slope:.4f} × {best_dataset} {best_intcpt:+.4f} m/s")

# ─────────────────────────────────────────────────────────────────────────────
# 4. MCP — APPLY TO FULL LTA PERIOD
# ─────────────────────────────────────────────────────────────────────────────
print()
print("="*60)
print("STEP 4 — LTA using best point")
print("="*60)

ref_lta = reanalysis_pts[(best_dataset, best_lat, best_lon)].copy()
ref_lta["lta_ws"] = np.clip(best_slope * ref_lta["ws"] + best_intcpt, 0, 50)
ref_lta["lta_wd"] = ref_lta["wd"]
ref_lta = ref_lta.dropna(subset=["lta_ws"])

lta_mean      = float(ref_lta["lta_ws"].mean())
scada_mean    = float(df_best_h["ws_turb"].mean())
print(f"  LTA mean wind speed          : {lta_mean:.2f} m/s")
print(f"  SCADA measurement period mean: {scada_mean:.2f} m/s")

# Sensitivity check: nearest point in selected dataset vs selected best-R point
same_ds_df = corr_df[corr_df["dataset"] == best_dataset].copy()
nearest_idx = int(same_ds_df["dist_km"].idxmin())
nearest_row = corr_df.loc[nearest_idx]
nearest_lat = float(nearest_row["lat"])
nearest_lon = float(nearest_row["lon"])
nearest_dist = float(nearest_row["dist_km"])
nearest_R = float(nearest_row["R"])
nearest_lta_mean = np.nan
lta_sens_delta = np.nan
if abs(nearest_lat - best_lat) < 1e-9 and abs(nearest_lon - best_lon) < 1e-9:
    nearest_lta_mean = lta_mean
    lta_sens_delta = 0.0
else:
    nearest_ref = reanalysis_pts[(best_dataset, nearest_lat, nearest_lon)].copy()
    nearest_lta_ws = np.clip(
        float(nearest_row["slope"]) * nearest_ref["ws"] + float(nearest_row["intercept"]),
        0, 50
    )
    nearest_lta_mean = float(nearest_lta_ws.mean())
    lta_sens_delta = lta_mean - nearest_lta_mean
print(f"  Sensitivity ({best_dataset} nearest @ {nearest_dist:.1f} km, "
      f"R={nearest_R:.4f}): ΔLTA = {lta_sens_delta:+.3f} m/s")

# ─────────────────────────────────────────────────────────────────────────────
# 5. STATISTICS
# ─────────────────────────────────────────────────────────────────────────────
print()
print("STEP 5 — Statistics")

# Weibull
ws_vals = ref_lta["lta_ws"].values
ws_vals = ws_vals[ws_vals > 0]
shape, loc, scale = weibull_min.fit(ws_vals, floc=0)
k_weib, A_weib = float(shape), float(scale)
print(f"  Weibull k={k_weib:.3f}  A={A_weib:.3f} m/s")

# Wind speed frequency 0.5 m/s bins
bins_05  = np.arange(0, 30.5, 0.5)
bin_cnts, _ = np.histogram(ws_vals, bins=bins_05)
bin_freq    = bin_cnts / bin_cnts.sum()
bin_lbls    = [f"({bins_05[i]:.1f},{bins_05[i+1]:.1f}]" for i in range(len(bins_05)-1)]

# Annual means
ref_lta["year"] = ref_lta.index.year
annual_ws       = ref_lta.groupby("year")["lta_ws"].mean()
annual_ws       = annual_ws[annual_ws.index <= int(LTA_END[:4])]
lta_mean_annual = float(annual_ws.mean())
print(f"  Annual means: {len(annual_ws)} years  LTA avg = {lta_mean_annual:.2f} m/s")

# Wind rose
wd_vals  = ref_lta["lta_wd"].dropna().values
ws_sync  = ref_lta.loc[ref_lta["lta_wd"].notna(), "lta_ws"].values
sec_idx  = ((wd_vals % 360) / 30).astype(int) % 12
sec_freq = np.array([(sec_idx == s).mean() for s in range(12)])
SEC_NAMES = [f"{i*30}–{(i+1)*30 if i < 11 else 360}" for i in range(12)]
SEC_EDGES_DEG = np.arange(0, 361, 30)

spd_bins = [0, 3, 6, 9, 12, 15, 18, 21, 30]
n_spd    = len(spd_bins) - 1
joint    = np.zeros((12, n_spd))
for s in range(12):
    mask = sec_idx == s
    if mask.sum() == 0:
        continue
    ws_s = ws_sync[mask]
    cnts = np.array([((ws_s > spd_bins[sp]) & (ws_s <= spd_bins[sp+1])).sum()
                     for sp in range(n_spd)], dtype=float)
    if cnts.sum() > 0:
        joint[s] = cnts / cnts.sum()

dom_sec  = int(np.argmax(sec_freq))
print(f"  Dominant sector: {SEC_NAMES[dom_sec]}°  ({sec_freq[dom_sec]*100:.1f}%)")
print("  Wind rose sectors (edge-aligned): 0–30, 30–60, ..., 330–360")

# ─────────────────────────────────────────────────────────────────────────────
# 6. FIGURES
# ─────────────────────────────────────────────────────────────────────────────
print()
print("STEP 6 — Generating figures")

def is_best_pt(dataset, lat, lon):
    return (dataset == best_dataset and
            abs(lat - best_lat) < 0.01 and
            abs(lon - best_lon) < 0.01)

# ── 6a. Map ──────────────────────────────────────────────────────────────────
# Map extent — tight around the wind farm and grid points
MAP_LAT_MIN, MAP_LAT_MAX = 48.30, 49.10
MAP_LON_MIN, MAP_LON_MAX = 4.80,  5.90

fig_map, ax_map = plt.subplots(figsize=(9, 7))
ax_map.set_xlim(MAP_LON_MIN, MAP_LON_MAX)
ax_map.set_ylim(MAP_LAT_MIN, MAP_LAT_MAX)

# Try to add terrain/satellite background with contextily
if CTX_OK:
    try:
        # Use OpenTopoMap for terrain context (no API key required)
        ctx.add_basemap(ax_map,
                        crs="EPSG:4326",
                        source=ctx.providers.OpenTopoMap,
                        zoom=10,
                        attribution=False,
                        zorder=0)
        print("  Terrain basemap added (OpenTopoMap).")
    except Exception as _e:
        print(f"  Terrain basemap unavailable ({_e}). Using plain background.")

# ERA5 points
for (lat, lon) in ERA5_POINTS:
    _best = is_best_pt("ERA5", lat, lon)
    ax_map.scatter(lon, lat, s=200, c="#0066CC",
                   marker="s", zorder=4,
                   edgecolors="darkred" if _best else "white",
                   linewidths=2.5 if _best else 0.8)
    dist  = haversine_km(SITE_LAT, SITE_LON, lat, lon)
    label = f"ERA5\n({lat:.2f}°N, {lon:.2f}°E)\nd={dist:.0f} km" + \
            ("\n★ SELECTED" if _best else "")
    ax_map.annotate(label, (lon, lat), textcoords="offset points",
                    xytext=(8, 6), fontsize=6.5,
                    color="darkred" if _best else "#003399",
                    fontweight="bold" if _best else "normal",
                    bbox=dict(boxstyle="round,pad=0.2", fc="white", alpha=0.7, lw=0),
                    zorder=5)

# MERRA-2 points
for (lat, lon) in MERRA_POINTS:
    _best = is_best_pt("MERRA-2", lat, lon)
    ax_map.scatter(lon, lat, s=200, c="#FF6600",
                   marker="^", zorder=4,
                   edgecolors="darkred" if _best else "white",
                   linewidths=2.5 if _best else 0.8)
    dist  = haversine_km(SITE_LAT, SITE_LON, lat, lon)
    label = f"MERRA-2\n({lat:.2f}°N, {lon:.3f}°E)\nd={dist:.0f} km" + \
            ("\n★ SELECTED" if _best else "")
    ax_map.annotate(label, (lon, lat), textcoords="offset points",
                    xytext=(8, 6), fontsize=6.5,
                    color="darkred" if _best else "#993300",
                    fontweight="bold" if _best else "normal",
                    bbox=dict(boxstyle="round,pad=0.2", fc="white", alpha=0.7, lw=0),
                    zorder=5)

# Individual turbines (tight cluster — shown as one group with labels)
for tid, (tlat, tlon) in TURBINE_COORDS.items():
    ax_map.scatter(tlon, tlat, s=80, c="red", marker="D", zorder=6,
                   edgecolors="darkred", linewidths=0.8)
    ax_map.annotate(tid, (tlon, tlat), textcoords="offset points",
                    xytext=(5, 4), fontsize=6.5, color="darkred",
                    fontweight="bold", zorder=7)

# Site reference label (centre of wind farm)
ax_map.annotate(f"{SITE_NAME}\nwind farm",
                (SITE_LON, SITE_LAT), textcoords="offset points",
                xytext=(-80, 18), fontsize=8.5, color="darkred",
                fontweight="bold",
                bbox=dict(boxstyle="round,pad=0.3", fc="white", alpha=0.85, lw=0.5, ec="darkred"),
                arrowprops=dict(arrowstyle="-", color="darkred", lw=0.8),
                zorder=8)

# Legend
legend_elems = [
    Line2D([0],[0], marker="s", linestyle="None", markerfacecolor="#0066CC",
           markersize=9,  markeredgecolor="white", label="ERA5 grid point (100 m)"),
    Line2D([0],[0], marker="^", linestyle="None", markerfacecolor="#FF6600",
           markersize=9,  markeredgecolor="white", label="MERRA-2 grid point (50 m)"),
    Line2D([0],[0], marker="D", linestyle="None", markerfacecolor="red",
           markersize=7,  markeredgecolor="darkred", label="Wind turbine (E1–E6)"),
    Line2D([0],[0], marker="s", linestyle="None", markerfacecolor="#0066CC",
           markersize=9,  markeredgecolor="darkred", markeredgewidth=2,
           label="Selected point (best R)"),
]
ax_map.legend(handles=legend_elems, loc="lower left", fontsize=8,
              framealpha=0.9, edgecolor="grey")
ax_map.set_xlabel("Longitude (°E)", fontsize=9)
ax_map.set_ylabel("Latitude (°N)",  fontsize=9)
ax_map.set_title(f"{SITE_NAME} Wind Farm — Site Location and Reanalysis Grid Points",
                 fontsize=10, fontweight="bold")
ax_map.tick_params(labelsize=8)
fig_map.tight_layout()
buf_map = io.BytesIO()
fig_map.savefig(buf_map, format="png", dpi=150, bbox_inches="tight")
buf_map.seek(0)
plt.close(fig_map)
print("  Map done.")

# ── 6b. Correlation scatter — all 8 points ───────────────────────────────────
fig_corr, axes = plt.subplots(2, 4, figsize=(14, 7))
axes_flat = axes.flatten()
for i, ((dataset, lat, lon), ref_df) in enumerate(reanalysis_pts.items()):
    ax  = axes_flat[i]
    mrg = overlap_cache[(dataset, lat, lon)]
    row = corr_df[(corr_df["dataset"]==dataset) &
                  (corr_df["lat"]==lat) &
                  (corr_df["lon"]==lon)].iloc[0]
    _best = is_best_pt(dataset, lat, lon)
    clr   = "#0066CC" if dataset == "ERA5" else "#FF6600"

    samp  = mrg.sample(min(3000, len(mrg)), random_state=42)
    ax.scatter(samp["ws_ref"], samp["ws_turb"], alpha=0.12, s=3, c=clr)
    x_r   = np.linspace(0, samp["ws_ref"].max(), 100)
    ax.plot(x_r, row["slope"]*x_r + row["intercept"],
            c=BLUE if dataset == "ERA5" else ORANGE, lw=1.8,
            label=f"R={row['R']:.4f}")
    ax.plot(x_r, x_r, "k--", lw=0.7, alpha=0.4, label="1:1")
    dist  = row["dist_km"]
    title = (f"{dataset}\n({lat:.2f}°N, {lon:.3f}°E)  d={dist:.0f} km" +
             ("\n[SELECTED ★]" if _best else ""))
    ax.set_title(title, fontsize=7.5,
                 fontweight="bold" if _best else "normal",
                 color="darkred" if _best else "black")
    ax.set_xlabel(f"{dataset} WS (m/s)", fontsize=7)
    ax.set_ylabel(f"{best_turbine} WS (m/s)", fontsize=7)
    ax.legend(fontsize=7, loc="upper left")
    ax.tick_params(labelsize=7)
    ax.grid(True, lw=0.4, alpha=0.4)
    if _best:
        for sp in ax.spines.values():
            sp.set_edgecolor("darkred")
            sp.set_linewidth(2.5)

fig_corr.suptitle(
    f"MCP Correlation — {best_turbine} vs 8 Reanalysis Points "
    f"(top row: ERA5, bottom row: MERRA-2)",
    fontsize=11, fontweight="bold")
fig_corr.tight_layout()
buf_corr = io.BytesIO()
fig_corr.savefig(buf_corr, format="png", dpi=130, bbox_inches="tight")
buf_corr.seek(0)
plt.close(fig_corr)
print("  Correlation scatter done.")

# ── 6c. Wind Rose ─────────────────────────────────────────────────────────────
fig_wr, ax_wr = plt.subplots(figsize=(6, 5.5), subplot_kw={"projection": "polar"})
ax_wr.set_theta_zero_location("N")
ax_wr.set_theta_direction(-1)
theta_edges = np.deg2rad(SEC_EDGES_DEG)
sector_w    = np.deg2rad(30)
bar_w       = sector_w * 0.85
colors_rose = plt.cm.Blues(np.linspace(0.3, 0.95, n_spd))
bottom      = np.zeros(12)
for sp in range(n_spd):
    h = joint[:, sp] * sec_freq * 100
    ax_wr.bar(theta_edges[:-1], h, width=bar_w, align="edge", bottom=bottom,
              color=colors_rose[sp], alpha=0.9, lw=0.3, edgecolor="white")
    bottom += h
ax_wr.set_xticks(np.deg2rad(np.arange(0, 360, 30)))
ax_wr.set_xticklabels(["N","NNE","ENE","E","ESE","SSE",
                       "S","SSW","WSW","W","WNW","NNW"], fontsize=8)
ax_wr.yaxis.set_tick_params(labelsize=7)
ax_wr.grid(True, lw=0.4, alpha=0.5)
spd_lbls = [f"{spd_bins[sp]}-{spd_bins[sp+1]} m/s" for sp in range(n_spd)]
handles  = [plt.Rectangle((0,0),1,1,fc=colors_rose[sp]) for sp in range(n_spd)]
ax_wr.legend(handles, spd_lbls, loc="lower left",
             bbox_to_anchor=(-0.25, -0.15), fontsize=7, ncol=2,
             title="Wind speed", title_fontsize=7)
ax_wr.set_title(f"Wind Rose — {SITE_NAME}\nMCP LTA ({LTA_START[:4]}–{LTA_END[:4]})",
                fontsize=9, fontweight="bold", pad=10)
fig_wr.tight_layout()
buf_wr = io.BytesIO()
fig_wr.savefig(buf_wr, format="png", dpi=150, bbox_inches="tight")
buf_wr.seek(0)
plt.close(fig_wr)
print("  Wind rose done.")

# ── 6d. Weibull ───────────────────────────────────────────────────────────────
fig_wb, ax_wb = plt.subplots(figsize=(7, 4))
x_ws     = np.linspace(0, 30, 300)
weib_pdf = weibull_min.pdf(x_ws, c=k_weib, loc=0, scale=A_weib)
bin_ctrs = (bins_05[:-1] + bins_05[1:]) / 2
ax_wb.bar(bin_ctrs, bin_freq, width=0.5, color=LB, alpha=0.6, label="LTA frequency")
ax_wb.plot(x_ws, weib_pdf * 0.5, color=BLUE, lw=2,
           label=f"Weibull fit  k={k_weib:.2f}, A={A_weib:.2f} m/s")
ax_wb.set_xlabel("Wind speed (m/s)", fontsize=9)
ax_wb.set_ylabel("Frequency",        fontsize=9)
ax_wb.set_xlim(0, 25)
ax_wb.set_title(f"Weibull Distribution — {SITE_NAME}", fontsize=9, fontweight="bold")
ax_wb.legend(fontsize=8)
ax_wb.grid(True, lw=0.4, alpha=0.5)
ax_wb.tick_params(labelsize=8)
fig_wb.tight_layout()
buf_wb = io.BytesIO()
fig_wb.savefig(buf_wb, format="png", dpi=150, bbox_inches="tight")
buf_wb.seek(0)
plt.close(fig_wb)
print("  Weibull done.")

# ── 6e. Annual means ──────────────────────────────────────────────────────────
fig_ann, ax_ann = plt.subplots(figsize=(9, 4))
yrs  = annual_ws.index.values
vals = annual_ws.values
clrs = [LB if v >= lta_mean_annual else "#A8C8E8" for v in vals]
bars = ax_ann.bar(yrs, vals, color=clrs, edgecolor="white", lw=0.5)
ax_ann.axhline(lta_mean_annual, color=BLUE, lw=1.5, ls="--",
               label=f"LTA mean = {lta_mean_annual:.2f} m/s")
ax_ann.set_xlabel("Year", fontsize=9)
ax_ann.set_ylabel("Mean wind speed (m/s)", fontsize=9)
ax_ann.set_title(f"Annual Mean Wind Speed (MCP LTA) — {SITE_NAME}",
                 fontsize=9, fontweight="bold")
ax_ann.set_xticks(yrs)
ax_ann.set_xticklabels(yrs, rotation=45, ha="right", fontsize=7)
ax_ann.tick_params(axis="y", labelsize=8)
ax_ann.legend(fontsize=8)
ax_ann.grid(True, axis="y", lw=0.4, alpha=0.5)
ax_ann.set_ylim(0, max(vals) * 1.15)
for bar, v in zip(bars, vals):
    ax_ann.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.04,
                f"{v:.2f}", ha="center", va="bottom", fontsize=6)
fig_ann.tight_layout()
buf_ann = io.BytesIO()
fig_ann.savefig(buf_ann, format="png", dpi=150, bbox_inches="tight")
buf_ann.seek(0)
plt.close(fig_ann)
print("  Annual chart done.")

# ── 6f. Power curve scatter — all turbines (Annex A) ─────────────────────────
fig_pc, axes_pc = plt.subplots(2, 3, figsize=(14, 9))
axes_pc_flat = axes_pc.flatten()
rng_pc = np.random.default_rng(42)

def _sample(mask, n_max):
    idx = np.where(mask)[0]
    if len(idx) > n_max:
        idx = rng_pc.choice(idx, n_max, replace=False)
    return idx

for i, tid in enumerate(TURBINE_IDS):
    ax   = axes_pc_flat[i]
    df_t  = turbine_data[tid]
    ws_r  = df_t["ws_raw"].values
    pw_v  = df_t["pw"].values
    flag   = df_t["flagged"].values
    curt_v = df_t["curtailed"].values

    has_pw      = np.isfinite(pw_v) & np.isfinite(ws_r)
    mask_blue   = has_pw & ~flag & ~curt_v   # clean
    mask_orange = has_pw & ~flag &  curt_v   # curtailment / derating
    mask_red    = has_pw &  flag             # excluded

    ib  = _sample(mask_blue,   6000)
    ior = _sample(mask_orange, 3000)
    ir  = _sample(mask_red,    2000)

    ax.scatter(ws_r[ib],  pw_v[ib],  s=1.5, alpha=0.12, c="#0066CC",
               label=f"Clean ({mask_blue.sum():,})",            rasterized=True)
    ax.scatter(ws_r[ior], pw_v[ior], s=2, alpha=0.30, c="#FF8800",
               label=f"Curtailment / derating ({mask_orange.sum():,})", rasterized=True)
    ax.scatter(ws_r[ir],  pw_v[ir],  s=5, alpha=0.55, c="red",
               label=f"Excluded ({mask_red.sum():,})",                  rasterized=True)
    ax.axhline(RATED_POWER, color="gray", lw=0.8, ls="--", alpha=0.6,
               label=f"Rated {RATED_POWER} kW")
    ax.set_xlim(-0.5, 25)
    ax.set_ylim(-300, RATED_POWER * 1.15)
    ax.set_xlabel("Wind speed (m/s)", fontsize=8)
    ax.set_ylabel("Active power (kW)", fontsize=8)
    ax.set_title(f"Turbine {tid}", fontsize=9, fontweight="bold")
    ax.legend(fontsize=6.5, loc="upper left", markerscale=3)
    ax.grid(True, lw=0.3, alpha=0.4)
    ax.tick_params(labelsize=7)

fig_pc.suptitle(
    f"{SITE_NAME} — Wind Speed vs Active Power (all turbines)\n"
    "Blue: clean  |  Orange: curtailment / derating  |  Red: excluded",
    fontsize=10, fontweight="bold"
)
fig_pc.tight_layout()
buf_pc = io.BytesIO()
fig_pc.savefig(buf_pc, format="png", dpi=130, bbox_inches="tight")
buf_pc.seek(0)
plt.close(fig_pc)
print("  Power curve scatter done.")

# ─────────────────────────────────────────────────────────────────────────────
# 6b. COVER PAGE IMAGE
# ─────────────────────────────────────────────────────────────────────────────
print("  Generating cover page image...")
_LOGO_PATH = os.path.join(BASE, "8p2 advisory white.png")
_FW, _FH   = 8.27, 11.69   # A4 in inches

fig_cov = plt.figure(figsize=(_FW, _FH))
ax_cov  = fig_cov.add_axes([0, 0, 1, 1])
ax_cov.set_xlim(0, _FW)
ax_cov.set_ylim(0, _FH)
ax_cov.set_facecolor("#001C3A")
ax_cov.axis("off")

# Sky gradient (subtle blue tint, lighter toward horizon)
_grad = np.linspace(1.0, 0.0, 512).reshape(-1, 1)
ax_cov.imshow(np.tile(_grad, (1, 2)),
              extent=[0, _FW, 0, _FH], aspect="auto",
              cmap=plt.cm.Blues, vmin=0.1, vmax=0.6,
              alpha=0.30, origin="upper", zorder=0)

from matplotlib.patches import Rectangle as _Rect, Polygon as _Poly, Circle as _Circ

# Wind turbine — tower centred at x=5.2
_tx, _ty0, _ty1 = 5.2, 0.0, 7.90    # tower x, base y (ground), hub y
_tw0, _tw1      = 0.20, 0.07         # half-width at base / top
ax_cov.add_patch(_Poly([
    (_tx - _tw0, _ty0), (_tx + _tw0, _ty0),
    (_tx + _tw1, _ty1), (_tx - _tw1, _ty1),
], closed=True, fc="#7AA5C8", ec="none", alpha=0.90, zorder=3))

# Nacelle
ax_cov.add_patch(_Poly([
    (_tx - 0.35, _ty1), (_tx + 0.65, _ty1),
    (_tx + 0.55, _ty1 + 0.22), (_tx - 0.28, _ty1 + 0.22),
], closed=True, fc="#9ABCD5", ec="none", alpha=0.90, zorder=4))

# Hub
_hx, _hy = _tx - 0.04, _ty1 + 0.11
ax_cov.add_patch(_Circ((_hx, _hy), 0.15, fc="#B8D4E8", ec="none", zorder=5))

# Three blades at 85°, 205°, 325° (slightly rotated from vertical for drama)
_R = 3.20
for _ang in [85, 205, 325]:
    _a   = np.radians(_ang)
    _bx  = _hx + _R * np.cos(_a)
    _by  = _hy + _R * np.sin(_a)
    _p90 = np.radians(_ang + 90)
    _bw  = 0.17    # root half-width
    ax_cov.add_patch(_Poly([
        (_hx + _bw * np.cos(_p90),         _hy + _bw * np.sin(_p90)),
        (_bx + 0.03 * np.cos(_p90),        _by + 0.03 * np.sin(_p90)),
        (_bx,                               _by),
        (_bx - 0.03 * np.cos(_p90),        _by - 0.03 * np.sin(_p90)),
        (_hx - _bw * np.cos(_p90),         _hy - _bw * np.sin(_p90)),
    ], closed=True, fc="#7AA5C8", ec="none", alpha=0.85, zorder=4))

# Text block (lower-left quadrant)
ax_cov.text(0.42, 3.95, SITE_NAME,
            color="white", fontsize=32, fontweight="bold",
            fontfamily="Open Sans", ha="left", va="bottom", zorder=10)
ax_cov.text(0.42, 3.15, "Long-Term Wind Distribution Analysis",
            color="white", fontsize=17, fontweight="bold",
            fontfamily="Open Sans", ha="left", va="bottom", zorder=10)
ax_cov.text(0.42, 2.50,
            f"{TURBINE_MODEL}  \u00d7{len(TURBINE_IDS)}  |  Meuse (55), Grand Est, France",
            color="white", fontsize=11, fontweight="bold",
            fontfamily="Open Sans", ha="left", va="bottom", zorder=10)
ax_cov.text(0.42, 1.90, f"{REPORT_RUN_DATE}  |  Prepared by 8.2 Advisory",
            color="white", fontsize=11, fontweight="bold",
            fontfamily="Open Sans", ha="left", va="bottom", zorder=10)

# 8.2 logo — top-left inset axes
try:
    _logo_img = plt.imread(_LOGO_PATH)
    _logo_ax  = fig_cov.add_axes([0.04, 0.90, 0.32, 0.08])
    _logo_ax.imshow(_logo_img)
    _logo_ax.axis("off")
except Exception as _le:
    print(f"  Cover logo not loaded: {_le}")

buf_cover = io.BytesIO()
fig_cov.savefig(buf_cover, format="png", dpi=150,
                bbox_inches="tight", pad_inches=0, facecolor="#001C3A")
buf_cover.seek(0)
plt.close(fig_cov)
print("  Cover page image done.")

# ─────────────────────────────────────────────────────────────────────────────
# 7. WORD DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────
print()
print("STEP 7 — Building Word document")

doc = Document()
# Set document default font
doc.styles["Normal"].font.name = "Open Sans"

# Main content section margins (TOC page 2 + body pages 3+)
doc.sections[0].top_margin    = Cm(2.0)
doc.sections[0].bottom_margin = Cm(2.0)
doc.sections[0].left_margin   = Cm(2.5)
doc.sections[0].right_margin  = Cm(2.5)

# ── Page 1: Cover ─────────────────────────────────────────────────────────────
p_cov = doc.add_paragraph()
p_cov.paragraph_format.space_before = Pt(0)
p_cov.paragraph_format.space_after  = Pt(0)
_run_cov = p_cov.add_run()
_run_cov.add_picture(buf_cover, width=Cm(21.0), height=Cm(29.7))

# Attach a paragraph-level sectPr to the cover paragraph.
# This defines the cover section (0 margins so image fills full A4 page).
# Everything after this paragraph uses the document-level sectPr (2/2.5 cm margins).
_pPr  = p_cov._p.get_or_add_pPr()
_sPr  = OxmlElement("w:sectPr")
_pgSz = OxmlElement("w:pgSz")
_pgSz.set(qn("w:w"), "11906")   # A4 width  in 20ths of a pt
_pgSz.set(qn("w:h"), "16838")   # A4 height in 20ths of a pt
_pgMar = OxmlElement("w:pgMar")
_pgMar.set(qn("w:top"),    "0")
_pgMar.set(qn("w:right"),  "0")
_pgMar.set(qn("w:bottom"), "0")
_pgMar.set(qn("w:left"),   "0")
_pgMar.set(qn("w:header"), "0")
_pgMar.set(qn("w:footer"), "0")
_pgMar.set(qn("w:gutter"), "0")
_sPr.append(_pgSz)
_sPr.append(_pgMar)
_pPr.append(_sPr)

# ── Page 2: Table of Contents ─────────────────────────────────────────────────
p_toc_h = doc.add_paragraph()
p_toc_h.paragraph_format.space_before = Pt(12)
p_toc_h.paragraph_format.space_after  = Pt(10)
_r_toc_h = p_toc_h.add_run("Table of Contents")
_r_toc_h.bold           = True
_r_toc_h.font.size      = Pt(16)
_r_toc_h.font.color.rgb = RGBColor.from_string("003366")
_r_toc_h.font.name      = "Open Sans"

# TOC field — Word auto-generates entries from Heading 1-3 styles on open/update
p_toc = doc.add_paragraph()
_r1 = p_toc.add_run()
_fc1 = OxmlElement("w:fldChar"); _fc1.set(qn("w:fldCharType"), "begin")
_r1._r.append(_fc1)

_r2 = p_toc.add_run()
_ins = OxmlElement("w:instrText")
_ins.set(qn("xml:space"), "preserve")
_ins.text = ' TOC \\o "1-3" \\h \\z \\u '
_r2._r.append(_ins)

_r3 = p_toc.add_run()
_fc2 = OxmlElement("w:fldChar"); _fc2.set(qn("w:fldCharType"), "separate")
_r3._r.append(_fc2)

_r4 = p_toc.add_run()
_r4.text    = "Right-click here and select \u2018Update Field\u2019 to generate the table of contents."
_r4.italic  = True
_r4.font.size = Pt(9)
_r4.font.name = "Open Sans"

_r5 = p_toc.add_run()
_fc3 = OxmlElement("w:fldChar"); _fc3.set(qn("w:fldCharType"), "end")
_r5._r.append(_fc3)

# ── UK-ordinal date helper (used in SCADA date references) ───────────────────
def _uk_date(dt):
    d = dt.day
    suffix = 'th' if 11 <= d <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(d % 10, 'th')
    return f"{d}{suffix} {dt.strftime('%B %Y')}"

_scada_start = _uk_date(df_best_h.index.min())
_scada_end   = _uk_date(df_best_h.index.max())


doc.add_page_break()

# ── 1. Site Description ───────────────────────────────────────────────────────
add_heading(doc, "1. Site Description", level=1)
add_body(doc,
    f"The {SITE_NAME} wind farm is located in the Meuse department (55), "
    f"Grand Est region, France. It consists of {len(TURBINE_IDS)} {TURBINE_MODEL} "
    f"wind turbines ({RATED_POWER} kW rated, Ø{ROTOR_DIAM} m rotor) arranged "
    f"along a roughly NE–SW axis. SCADA data is available from "
    f"{_scada_start} to {_scada_end}."
)
add_body(doc,
    f"Terminology note: throughout this report, 'LTA' denotes the Long-Term Average "
    f"wind speed at hub height, derived by applying the MCP regression to the "
    f"{int(LTA_END[:4]) - int(LTA_START[:4]) + 1}-year reanalysis reference period "
    f"({LTA_START[:4]}–{LTA_END[:4]}). This should not be confused with "
    f"'Long-Term Assessment' (also abbreviated LTA in other technical contexts)."
)

add_heading(doc, "1.1  Hub Height Note", level=2)
p_hub = doc.add_paragraph()
p_hub.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r_hub = p_hub.add_run(
    f"Hub height confirmed as {HUB_HEIGHT} m from the project database (1&3.xlsx). "
    f"ERA5 data is provided at 100 m (u100/v100) and therefore requires no height "
    f"correction (factor = 1.000). MERRA-2 data is at 50 m (U50M/V50M) and is "
    f"corrected to 100 m using the power law with α = {ALPHA_SHEAR} "
    f"(factor = {MERRA_CORR:.4f})."
)
r_hub.font.size = Pt(10)
r_hub.font.name = "Open Sans"
p_hub.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

add_heading(doc, "1.2  Turbine Coordinates", level=2)
tbl_gps = doc.add_table(rows=len(TURBINE_IDS)+1, cols=4)
tbl_gps.style = "Table Grid"
table_header(tbl_gps, ["Turbine", "Latitude (°N)", "Longitude (°E)", "Distance from centroid (km)"])
for ri, tid in enumerate(TURBINE_IDS):
    tlat, tlon = TURBINE_COORDS[tid]
    dist_cent  = haversine_km(SITE_LAT, SITE_LON, tlat, tlon)
    row = tbl_gps.rows[ri+1]
    for ci, val in enumerate([tid, f"{tlat:.6f}", f"{tlon:.6f}", f"{dist_cent:.3f}"]):
        row.cells[ci].text = val
        row.cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
        if ri % 2 == 0:
            set_cell_bg(row.cells[ci], "EBF2FA")

doc.add_paragraph()

add_heading(doc, "1.3  Site Location and Reanalysis Grid Points", level=2)
add_body(doc,
    "The following map shows the six turbine positions and all 8 reanalysis grid "
    "points (4 ERA5 at 0.25° resolution, 4 MERRA-2 at ~0.5° resolution). "
    "The selected point for the MCP analysis is indicated by ★."
)
doc.add_picture(buf_map, width=Inches(5.8))
add_caption(doc,
    f"Figure 1 — Site location and surrounding reanalysis grid points. "
    f"Blue squares: ERA5 (100 m wind). Orange triangles: MERRA-2 (50 m wind). "
    f"Red diamonds: individual turbines (E1–E6). "
    f"★ Selected point: {best_dataset} ({best_lat}°N, {best_lon}°E)."
)

doc.add_paragraph()

add_heading(doc, "1.4  Reanalysis Points Summary", level=2)
add_body(doc,
    "The table below lists all 8 candidate reanalysis grid points with their "
    "distance from the site centroid. All points are height-corrected to the "
    f"confirmed hub height of {HUB_HEIGHT} m using the power law (α = {ALPHA_SHEAR})."
)
tbl_pts = doc.add_table(rows=len(corr_df)+1, cols=6)
tbl_pts.style = "Table Grid"
table_header(tbl_pts, ["Dataset", "Latitude (°N)", "Longitude (°E)",
                        "Distance (km)", "Source height (m)", "Correction factor"])
for ri, (_, rw) in enumerate(corr_df.iterrows()):
    row    = tbl_pts.rows[ri+1]
    _best  = is_best_pt(rw["dataset"], rw["lat"], rw["lon"])
    h_src  = ERA5_HEIGHT if rw["dataset"] == "ERA5" else MERRA_HEIGHT
    corr_f = ERA5_CORR   if rw["dataset"] == "ERA5" else MERRA_CORR
    vals   = [rw["dataset"] + (" [SELECTED]" if _best else ""),
              f"{rw['lat']:.2f}", f"{rw['lon']:.3f}",
              f"{rw['dist_km']:.1f}", str(h_src), f"{corr_f:.4f}"]
    for ci, val in enumerate(vals):
        row.cells[ci].text = val
        row.cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
        if _best:
            set_cell_bg(row.cells[ci], "FFF0E0")
            row.cells[ci].paragraphs[0].runs[0].bold = True
        elif ri % 2 == 0:
            set_cell_bg(row.cells[ci], "EBF2FA")

doc.add_page_break()

# ── 2. Methodology ────────────────────────────────────────────────────────────
add_heading(doc, "2. Methodology", level=1)
add_body(doc,
    "The analysis follows a Measure-Correlate-Predict (MCP) approach to extend "
    "the short-term SCADA measurement record to a long-term wind distribution "
    f"representative of the full reference period ({LTA_START[:4]}–{LTA_END[:4]}). "
    "The following steps were applied:"
)

add_heading(doc, "2.1  Step 1 — Turbine Selection (Conservative Approach)", level=2)
add_body(doc,
    "The mean wind speed is computed for each of the six turbines from the available "
    "SCADA data. The turbine with the highest mean wind speed is selected as the "
    "reference for the MCP correlation. This is the most conservative choice for a "
    "structural lifetime assessment: it represents the machine experiencing the "
    "greatest cumulative wind exposure, which produces the highest fatigue load estimates."
)

add_heading(doc, "2.2  Step 2 — SCADA Data Cleaning", level=2)
add_body(doc,
    "The selected turbine's wind speed data is cleaned before correlation to remove "
    "physically invalid or instrumentally suspect records:"
)
for _st in [
    "Physical outliers: records with wind speed < 0 m/s or > 50 m/s are removed.",
    f"Frozen sensor detection: runs of ≥ {FROZEN_MIN_RECORDS} consecutive identical "
    f"wind speed values > 0.1 m/s (non-zero) in the native 10-minute SCADA data "
    f"(≥ {FROZEN_MIN_RECORDS * 10} minutes) are flagged and removed. "
    "Exactly-zero and near-zero records (≤ 0.1 m/s) are excluded from this check "
    "as sustained calms cannot be distinguished from a genuinely still atmosphere. "
    "The criterion detects a stuck or frozen nacelle anemometer reading any "
    "non-trivial wind speed.",
    f"Zero wind-speed inconsistency: samples with windSpeedAvg = 0 m/s are retained "
    f"unless inconsistent with turbine production. Records with windSpeedAvg = 0 m/s "
    f"and powerAvg > {ZERO_WS_PMIN_KW:.1f} kW ({100 * ZERO_WS_PMIN_KW / RATED_POWER:.1f}% "
    f"of rated power) are removed as physically implausible.",
]:
    _p = doc.add_paragraph(style="List Bullet")
    _r = _p.add_run(_st)
    _r.font.size = Pt(10)
    _r.font.name = "Open Sans"

add_heading(doc, "2.3  Step 3 — Correlation with Reanalysis (Wind Speed Only)", level=2)
add_body(doc,
    "The native SCADA timestep is 10 minutes. Frozen-sensor detection (Step 2 above) "
    "is applied at that resolution, before any resampling. After cleaning, the "
    "10-minute wind speed values are aggregated to hourly means; only hours with "
    "at least one valid 10-minute record contribute to the MCP correlation. "
    "This hourly cleaned series from the selected turbine is then correlated "
    "with each of 8 candidate reanalysis grid points: 4 ERA5 nodes (0.25° resolution, "
    "100 m wind level) and 4 MERRA-2 nodes (~0.5° resolution, 50 m wind level). "
    "The reanalysis data is also used at hourly resolution. "
    "Height correction to hub height is applied using the power law "
    f"(α = {ALPHA_SHEAR}) where required."
)
add_body(doc,
    "Only wind speed (scalar magnitude, m/s) is regressed. Wind direction is not "
    "correlated — it is taken directly from the selected reanalysis point for the "
    "full LTA period. Reanalysis direction data is physically reliable and free of "
    "the nacelle-wake bias that affects anemometer measurements."
)
add_body(doc,
    "A linear regression SCADA = a × Reanalysis + b is fitted for each of the "
    "8 candidate points. The Pearson correlation coefficient R is computed for "
    "each pair; the point with the highest R is selected as the reference for "
    "the long-term prediction."
)

add_heading(doc, "2.4  Wind Measurement Source — Nacelle Anemometry", level=2)
add_body(doc,
    "The wind speed data used in this analysis comes from the SCADA operational "
    "records of each turbine. Wind speed in SCADA systems is characteristically "
    "recorded by a nacelle-mounted anemometer positioned immediately behind the "
    "rotor hub — not by a free-standing meteorological mast measuring undisturbed "
    "free-stream wind upstream of the rotor. This assumption is based on the "
    "standard architecture of wind turbine SCADA systems; if a dedicated met mast "
    "or remote sensing device (LiDAR, SoDAR) is available at this site, its data "
    "should be preferred, as it provides free-stream wind speed measurements with "
    "typical correlation values R > 0.92 against reanalysis."
)
add_body(doc,
    "A nacelle transfer function (NTF), as defined in IEC 61400-12-2, converts "
    "the wake-affected nacelle wind speed measurement to an equivalent free-stream "
    "wind speed at hub height. The NTF is derived empirically by comparing nacelle "
    "readings against a reference met mast during a calibration campaign. If the "
    "turbine manufacturer (REpower/Senvion) provides an NTF for the MM92 model, "
    "applying it before MCP correlation would reduce the systematic bias and improve "
    "the correlation coefficient. No NTF has been applied in this analysis."
)
add_body(doc,
    "The wind speed vs power scatter plots in Annex A provide a visual verification "
    "that the SCADA data follows the characteristic shape of a pitch-regulated "
    "turbine power curve, consistent with nacelle anemometry."
)

add_heading(doc, "2.5  Note on Expected Correlation Values", level=2)
add_body(doc,
    "Correlations between nacelle anemometer data and gridded reanalysis are "
    "typically in the range R = 0.80–0.90, which is lower than the R ≥ 0.92 "
    "achieved with a free-standing met mast. This is expected and does not affect "
    "the validity of the MCP result. Three effects reduce R:"
)
for _r in [
    "Nacelle wake: the anemometer is mounted behind the rotor hub and operates in "
    "the disturbed wake of the rotor. When generating power, the rotor extracts "
    "energy from the wind, reducing the measured speed below the free-stream value. "
    "This effect varies with wind speed and yaw error, adding scatter to the "
    "correlation.",
    "Yaw misalignment: small tracking errors between nacelle heading and wind "
    "direction mean the anemometer is not always aligned with the flow.",
    "Spatial scale mismatch: ERA5 at ~28 km grid spacing and MERRA-2 at ~50 km "
    "cannot resolve local terrain channelling or turbulence, introducing "
    "irreducible scatter between point measurements and grid-cell averages.",
]:
    _p = doc.add_paragraph(style="List Bullet")
    _rr = _p.add_run(_r)
    _rr.font.size = Pt(10)
    _rr.font.name = "Open Sans"
add_body(doc,
    "Despite lower R values, the MCP regression correctly captures the systematic "
    "wind-speed relationship between the reanalysis source and the site. The "
    "long-term mean and distribution derived by MCP are not materially sensitive "
    "to R in the 0.80–0.90 range when the overlap period is sufficiently long "
    "(several years of concurrent data)."
)

add_heading(doc, "2.6  Step 4 — Long-Term Prediction", level=2)
add_body(doc,
    "The selected MCP regression equation is applied to the full reanalysis record "
    f"({LTA_START[:4]}–{LTA_END[:4]}) to generate a synthetic long-term wind speed "
    "time series at hub height. The Weibull distribution parameters, annual mean "
    "wind speeds, and wind rose are all derived from this synthetic series."
)

doc.add_page_break()

# ── 3. Data Quality & Turbine Selection ──────────────────────────────────────
add_heading(doc, "3. SCADA Data Quality and Turbine Selection", level=1)
add_body(doc,
    f"All 6 SCADA wind speed time series were cleaned as follows before MCP correlation:"
)
steps = [
    f"Physical outliers: records with wind speed < 0 m/s or > 50 m/s were removed.",
    f"Frozen sensor: records where the wind speed exceeded 0.1 m/s and remained "
    f"identical for ≥ {FROZEN_MIN_RECORDS} consecutive 10-minute records "
    f"(≥ {FROZEN_MIN_RECORDS * 10} min) were removed. "
    f"Near-zero records (≤ 0.1 m/s) are excluded from the check as sustained "
    f"calms cannot be unambiguously classified as sensor faults.",
    f"Zero wind-speed inconsistency: records with windSpeedAvg = 0 m/s were retained "
    f"unless powerAvg exceeded {ZERO_WS_PMIN_KW:.1f} kW "
    f"({100 * ZERO_WS_PMIN_KW / RATED_POWER:.1f}% of rated power), in which case "
    f"the wind-speed value was removed as physically implausible.",
]
for st in steps:
    _p = doc.add_paragraph(style="List Bullet")
    _r = _p.add_run(st)
    _r.font.size = Pt(10)
    _r.font.name = "Open Sans"

doc.add_paragraph()

add_heading(doc, "3.1  Per-Turbine Statistics After Cleaning", level=2)
tbl_st = doc.add_table(rows=len(TURBINE_IDS)+1, cols=6)
tbl_st.style = "Table Grid"
table_header(tbl_st, ["Turbine", "Total records", "Frozen removed",
                       "Zero WS @ power removed", "Clean availability (%)",
                       "Mean wind speed (m/s)"])
for ri, tid in enumerate(TURBINE_IDS):
    st   = stats_df.loc[tid]
    row  = tbl_st.rows[ri+1]
    isel = (tid == best_turbine)
    vals = [tid + (" [SELECTED]" if isel else ""),
            f"{st['n_total']:,}", f"{st['n_frozen']:,}",
            f"{st['n_zero_ws_power']:,}",
            f"{st['avail_pct']:.1f}%",
            f"{st['mean_ws']:.2f}"]
    for ci, val in enumerate(vals):
        row.cells[ci].text = val
        row.cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
        if isel:
            set_cell_bg(row.cells[ci], "FFF0E0")
            row.cells[ci].paragraphs[0].runs[0].bold = True
        elif ri % 2 == 0:
            set_cell_bg(row.cells[ci], "EBF2FA")

doc.add_paragraph()
add_body(doc,
    f"Turbine {best_turbine} has the highest clean mean wind speed "
    f"({best_ws_mean:.2f} m/s) and is therefore selected for the MCP analysis. "
    f"This represents the most conservative approach for a wind-load lifetime "
    f"assessment, as it uses the turbine experiencing the greatest wind exposure."
)

doc.add_page_break()

# ── 3. MCP Correlation ────────────────────────────────────────────────────────
add_heading(doc, "4. MCP Correlation Analysis", level=1)
add_body(doc,
    f"The cleaned SCADA wind speed data from turbine {best_turbine} (native 10-minute "
    f"records aggregated to hourly means for MCP) was correlated "
    f"with all 8 reanalysis grid points over the concurrent measurement period. "
    f"A linear regression (SCADA = a × Reanalysis + b) is applied. "
    f"The point with the highest Pearson correlation R is selected for the "
    f"long-term prediction."
)

add_heading(doc, "4.1  Correlation Results — All 8 Points", level=2)
tbl_corr = doc.add_table(rows=len(corr_df)+1, cols=7)
tbl_corr.style = "Table Grid"
table_header(tbl_corr, ["Dataset", "Lat (°N)", "Lon (°E)",
                         "Dist (km)", "n overlap", "R", "Selected"])
for ri, (_, rw) in enumerate(corr_df.iterrows()):
    row   = tbl_corr.rows[ri+1]
    _best = is_best_pt(rw["dataset"], rw["lat"], rw["lon"])
    vals  = [rw["dataset"], f"{rw['lat']:.2f}", f"{rw['lon']:.3f}",
             f"{rw['dist_km']:.1f}", f"{rw['n_overlap']:,}",
             f"{rw['R']:.4f}", "YES ★" if _best else ""]
    for ci, val in enumerate(vals):
        row.cells[ci].text = val
        row.cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
        if _best:
            set_cell_bg(row.cells[ci], "FFF0E0")
            row.cells[ci].paragraphs[0].runs[0].bold = True
        elif ri % 2 == 0:
            set_cell_bg(row.cells[ci], "EBF2FA")

doc.add_paragraph()
add_body(doc,
    f"The best correlation is with {best_dataset} at ({best_lat}°N, {best_lon}°E), "
    f"R = {best_R:.4f}. MCP regression: "
    f"SCADA ({best_turbine}) = {best_slope:.4f} × {best_dataset} "
    f"{_sign_intcpt}{abs(best_intcpt):.4f} m/s."
)
if np.isfinite(lta_sens_delta):
    _sens_sign = "+" if lta_sens_delta >= 0 else ""
    add_body(doc,
        f"Correlation differences are marginal among nearby {best_dataset} points. "
        f"For transparency, using the nearest {best_dataset} point "
        f"({nearest_lat:.2f}°N, {nearest_lon:.3f}°E; distance {nearest_dist:.1f} km; "
        f"R = {nearest_R:.4f}) changes the LTA mean by {_sens_sign}{lta_sens_delta:.3f} m/s "
        f"relative to the selected best-R point."
    )
doc.add_paragraph()

_dir = "higher" if scada_mean > lta_mean else "lower"
_over = "above" if scada_mean > lta_mean else "below"
_bias = "overestimate" if scada_mean > lta_mean else "underestimate"
add_body(doc,
    f"Note on measurement period vs long-term average: the clean SCADA mean wind "
    f"speed for turbine {best_turbine} over the measurement period "
    f"({_scada_start} to {_scada_end}) is "
    f"{scada_mean:.2f} m/s, which is {_dir} than the long-term average of "
    f"{lta_mean:.2f} m/s over {LTA_START[:4]}–{LTA_END[:4]}. The measurement "
    f"period coincided with {_over}-average wind conditions, which is a normal "
    f"consequence of inter-annual wind variability. Using the raw SCADA mean "
    f"alone (without MCP) would {_bias} the long-term site wind resource. "
    f"The MCP method corrects for this by anchoring the prediction to the full "
    f"{int(LTA_END[:4]) - int(LTA_START[:4]) + 1}-year reanalysis climatology."
)
doc.add_paragraph()

add_heading(doc, "4.2  Scatter Plots — All 8 Points", level=2)
doc.add_picture(buf_corr, width=Inches(6.5))
add_caption(doc,
    f"Figure 2 — MCP scatter plots for {best_turbine} versus all 8 reanalysis points. "
    f"Top row: ERA5 (100 m corrected to {HUB_HEIGHT} m). "
    f"Bottom row: MERRA-2 (50 m corrected to {HUB_HEIGHT} m). "
    f"Red-bordered panel: selected point "
    f"({best_dataset}, {best_lat}°N, {best_lon}°E, R = {best_R:.4f})."
)

doc.add_page_break()

# ── 4. Wind Direction ─────────────────────────────────────────────────────────
add_heading(doc, "5. Wind Direction Data", level=1)
add_body(doc,
    f"Wind direction is taken from the selected reanalysis point "
    f"({best_dataset}, {best_lat}°N, {best_lon}°E) after MCP prediction over "
    f"{LTA_START[:4]}–{LTA_END[:4]}. "
    f"Directions are meteorological (from which the wind blows)."
)

add_heading(doc, "5.1  Wind Rose", level=2)
doc.add_picture(buf_wr, width=Inches(5.0))
add_caption(doc,
    f"Figure 3 — Wind rose based on MCP LTA ({LTA_START[:4]}–{LTA_END[:4]}). "
    f"12 sectors of 30° each. Colours by wind speed class."
)
doc.add_paragraph()

add_heading(doc, "5.2  Frequency by 30° Sector", level=2)
tbl_dir = doc.add_table(rows=13, cols=3)
tbl_dir.style = "Table Grid"
table_header(tbl_dir, ["Sector (°)", "Mid-point (°)", "Frequency (%)"])
for s in range(12):
    row = tbl_dir.rows[s+1]
    for ci, val in enumerate([SEC_NAMES[s], str(s*30+15), f"{sec_freq[s]*100:.2f}%"]):
        row.cells[ci].text = val
        row.cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
        if s % 2 == 0:
            set_cell_bg(row.cells[ci], "EBF2FA")

doc.add_page_break()

# ── 5. Wind Speed Distribution ────────────────────────────────────────────────
add_heading(doc, "6. Wind Speed Distribution", level=1)

add_heading(doc, "6.1  Weibull Distribution", level=2)
doc.add_picture(buf_wb, width=Inches(6.2))
add_caption(doc,
    f"Figure 4 — Weibull distribution fit to MCP LTA "
    f"({LTA_START[:4]}–{LTA_END[:4]}). "
    f"Shape k = {k_weib:.3f},  scale A = {A_weib:.3f} m/s.  "
    f"LTA mean = {lta_mean:.2f} m/s."
)
doc.add_paragraph()

add_heading(doc, "6.2  Wind Speed Frequency per 0.5 m/s Bin", level=2)
tbl_ws   = doc.add_table(rows=len(bin_lbls)+1, cols=2)
tbl_ws.style = "Table Grid"
table_header(tbl_ws, ["Bin (m/s)", "Frequency (%)"])
for i in range(len(bin_lbls)):
    row = tbl_ws.rows[i+1]
    row.cells[0].text = bin_lbls[i]
    row.cells[1].text = f"{bin_freq[i]*100:.4f}%"
    for ci in range(2):
        row.cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
        if (i + 1) % 2 == 0:
            set_cell_bg(row.cells[ci], "EBF2FA")

doc.add_paragraph()
tail_prob = float((ref_lta["lta_ws"] > 16.0).mean() * 100.0)
add_body(doc,
    f"All 0.5 m/s bins from 0 to 30 m/s are reported in the table above. "
    f"The tail probability above 16 m/s is {tail_prob:.3f}%."
)

add_heading(doc, "6.3  Annual Mean Wind Speed", level=2)
doc.add_picture(buf_ann, width=Inches(6.5))
add_caption(doc,
    f"Figure 5 — Annual mean wind speed (MCP LTA, {LTA_START[:4]}–{LTA_END[:4]}). "
    f"LTA average = {lta_mean_annual:.2f} m/s. "
    f"SCADA period mean ({best_turbine}) = {scada_mean:.2f} m/s."
)
doc.add_paragraph()

tbl_ann = doc.add_table(rows=len(annual_ws)+2, cols=2)
tbl_ann.style = "Table Grid"
table_header(tbl_ann, ["Year", "Annual mean wind speed (m/s)"])
for ri, (yr, ws_yr) in enumerate(annual_ws.items()):
    row = tbl_ann.rows[ri+1]
    for ci, val in enumerate([str(yr), f"{ws_yr:.2f}"]):
        row.cells[ci].text = val
        row.cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
        if ri % 2 == 0:
            set_cell_bg(row.cells[ci], "EBF2FA")
# LTA row
row_lta = tbl_ann.rows[-1]
for ci, val in enumerate([f"LTA {LTA_START[:4]}–{LTA_END[:4]}", f"{lta_mean_annual:.2f}"]):
    row_lta.cells[ci].text = val
    set_cell_bg(row_lta.cells[ci], "003366")
    row_lta.cells[ci].paragraphs[0].runs[0].bold           = True
    row_lta.cells[ci].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    row_lta.cells[ci].paragraphs[0].runs[0].font.size      = Pt(9)

doc.add_page_break()

# ── 6. Summary ────────────────────────────────────────────────────────────────
add_heading(doc, "7. Summary", level=1)
add_body(doc,
    f"This report presents a long-term wind distribution analysis for the "
    f"{SITE_NAME} wind farm, comprising {len(TURBINE_IDS)} {TURBINE_MODEL} turbines "
    f"({RATED_POWER} kW rated, {ROTOR_DIAM} m rotor diameter, {HUB_HEIGHT} m hub height) "
    f"located in the Meuse département (55), Grand Est, France. "
    f"SCADA ten-minute records were extracted for all six turbines and cleaned of "
    f"physical outliers and frozen-sensor events. Zero wind-speed values were retained "
    f"unless inconsistent with production (windSpeedAvg = 0 m/s with powerAvg > "
    f"{ZERO_WS_PMIN_KW:.1f} kW), in which case they were removed. "
    f"Curtailment / derating periods were "
    f"identified using the REpower MM92 manufacturer power curve as a fixed external "
    f"reference; wind speed was retained for MCP at these records, as the nacelle "
    f"anemometer operates independently of turbine operating state. "
    f"The turbine with the highest clean mean wind speed — "
    f"{best_turbine} at {best_ws_mean:.2f} m/s — was selected for the "
    f"Measure-Correlate-Predict (MCP) regression as a conservative approach."
)
add_body(doc,
    f"MCP linear regression was performed against eight candidate reanalysis grid "
    f"points (four ERA5 at 100 m, four MERRA-2 at 50 m). The strongest correlation "
    f"(Pearson R = {best_R:.4f}) was found at the {best_dataset} grid point "
    f"({best_lat}\u00b0N, {best_lon}\u00b0E), yielding the equation: "
    f"SCADA\u2009({best_turbine}) = {best_slope:.4f}\u2009\u00d7\u2009{best_dataset} "
    f"{best_intcpt:+.4f}\u2009m/s. "
    f"This equation was applied to the full {int(LTA_END[:4])-int(LTA_START[:4])+1}-year "
    f"reanalysis record ({LTA_START[:4]}\u2013{LTA_END[:4]}) to derive the long-term "
    f"wind distribution at hub height."
)
add_body(doc,
    f"The long-term mean wind speed at {HUB_HEIGHT} m hub height is {lta_mean:.2f} m/s, "
    f"compared with a SCADA measurement-period mean of {scada_mean:.2f} m/s for "
    f"{best_turbine}. The wind speed distribution is well described by a Weibull "
    f"function with shape parameter k\u2009=\u2009{k_weib:.3f} and scale parameter "
    f"A\u2009=\u2009{A_weib:.3f}\u2009m/s. The dominant wind direction is the "
    f"{SEC_NAMES[dom_sec]}\u00b0 sector ({sec_freq[dom_sec]*100:.1f}\u2009% frequency), "
    f"consistent with the prevailing south-westerly flow typical of north-eastern France."
)

doc.add_page_break()

# ── Annex A. Power Curve Scatter Plots ────────────────────────────────────────
add_heading(doc, "Annex A — Wind Speed vs Power Curves", level=1)
add_body(doc,
    "The following scatter plots show the measured wind speed versus active power "
    "for all six turbines, using the full SCADA record. The plots serve as a visual "
    "quality check of the data before and after cleaning."
)
for _itm in [
    "Blue points: clean data retained in the analysis — normal turbine operation "
    "with valid wind speed and power.",
    f"Orange points: probable curtailment or derating — active power is below "
    f"{int(MM92_CURT_THRESHOLD * 100)} % of the REpower MM92 manufacturer reference "
    f"power curve at the measured wind speed (IEC standard air density 1.225 kg/m\u00b3). "
    f"Using a fixed external reference curve avoids the circular-reference problem "
    f"of data-derived thresholds. Wind speed is retained for MCP — the nacelle "
    f"anemometer reads independently of turbine operating state.",
    "Red points: records excluded during cleaning — physical outliers (wind speed "
    "outside 0–50 m/s) or frozen-sensor detections (≥ 6 consecutive identical "
    "non-zero values).",
    "Downtime and low-production periods (turbine stopped or ramping at low wind "
    "speed) appear as blue points near zero power, which is normal behaviour.",
]:
    _p = doc.add_paragraph(style="List Bullet")
    _ra = _p.add_run(_itm)
    _ra.font.size = Pt(10)
    _ra.font.name = "Open Sans"
add_body(doc,
    "The characteristic S-shaped power curve (cut-in around 3 m/s, rated power "
    f"at approximately 12–14 m/s, cut-out above ~25 m/s) is consistent with the "
    f"{TURBINE_MODEL} specification and confirms that the SCADA wind speed data "
    "originates from nacelle-mounted anemometers. The orange curtailment band "
    "is clearly visible on all turbines and reflects periods of grid or noise "
    "constraint, not instrument error. "
    "The scatter at intermediate wind speeds (5–12 m/s) is typical of nacelle "
    "anemometry, as described in Section 2.4."
)
add_body(doc,
    "Note on curtailment and the long-term wind distribution: curtailment events "
    "do not affect the LTA result. The nacelle anemometer measures the approaching "
    "wind speed independently of the turbine operating state — whether the turbine "
    "is producing full power, curtailed, or stopped. The MCP correlation is "
    "performed on wind speed alone; the long-term mean wind speed, Weibull "
    "distribution, and wind rose derived from the MCP prediction are therefore "
    "unaffected by periods of curtailment or derating. Curtailment detection uses "
    f"the REpower MM92 manufacturer reference power curve as a fixed external "
    f"threshold ({int(MM92_CURT_THRESHOLD*100)} % of rated reference) rather than "
    f"a data-derived percentile, so sustained power caps at any level are reliably "
    f"identified."
)
doc.add_picture(buf_pc, width=Inches(6.5))
add_caption(doc,
    f"Figure A.1 — Wind speed vs active power for all six {SITE_NAME} turbines. "
    f"Blue: clean data (normal operation). "
    f"Orange: curtailed records (wind speed retained for MCP). "
    f"Red: excluded data (physical outlier or frozen anemometer). "
    f"Dashed grey line: rated power ({RATED_POWER} kW). "
    f"SCADA period {_scada_start} to {_scada_end}."
)

doc.save(OUT_DOCX)
print(f"\nDocument saved: {OUT_DOCX}")
print(f"\nSite centroid used: {SITE_LAT:.4f}°N, {SITE_LON:.4f}°E")
